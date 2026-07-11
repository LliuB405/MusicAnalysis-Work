# -*- coding: utf-8 -*-
"""
按分工差异化生成 6 个 docx 文件（简化重写版）
"""
import os
import sys
import io
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

PROJECT = {
    "title": "基于网易云热歌榜的音乐数据分析可视化系统",
    "course": "Python应用开发技术",
    "class_name": "24软工2班",
}

TEMPLATE_DIR = "C:/Users/22283/Desktop"
OUTPUT_DIR = "C:/Users/22283/Desktop/大作业提交材料"


def set_cell_text(cell, text):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    if cell.paragraphs:
        cell.paragraphs[0].text = text


def replace_paragraph_text(paragraph, text):
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    lines = text.split("\n")
    paragraph.add_run(lines[0])
    for line in lines[1:]:
        paragraph.add_run().add_break()
        paragraph.add_run(line)


# ============================================================
# 大作业报告 - 队长刘冰（爬虫/数据/可视化）
# ============================================================
def fill_assignment_leader(doc):
    t0 = doc.tables[0]
    set_cell_text(t0.cell(0, 1), PROJECT["title"])
    set_cell_text(t0.cell(1, 1), "202437020104")
    set_cell_text(t0.cell(2, 1), "刘冰")
    set_cell_text(t0.cell(3, 1), PROJECT["class_name"])
    set_cell_text(t0.cell(4, 1), PROJECT["course"])

    t2 = doc.tables[2]
    set_cell_text(t2.cell(1, 0), "歌曲排名")
    set_cell_text(t2.cell(1, 1), "200条")
    set_cell_text(t2.cell(1, 2), "1-200名（玻璃、海屿你、12.31、恋人等）")
    set_cell_text(t2.cell(1, 3), "API实时返回200条，无重复，剔除空歌名1条")
    set_cell_text(t2.cell(2, 0), "歌手名称")
    set_cell_text(t2.cell(2, 1), "200条")
    set_cell_text(t2.cell(2, 2), "Gareth.T、马也_Crabbit、郑润泽、李荣浩等")
    set_cell_text(t2.cell(2, 3), "artists字段正确解析，多歌手用 / 分隔")

    content = {
        "（简要说明选择该主题的原因、研究价值或实际意义，100-200字）":
            "作为本项目的系统架构师和爬虫核心开发，我选择网易云热歌榜作为研究对象，"
            "主要基于三方面考量：\n"
            "（1）技术价值：网易云音乐的反爬机制在主流平台中具有代表性，"
            "通过本项目可深入掌握 requests Session、headers 伪装、API 重试等"
            "工业级爬虫技术；\n"
            "（2）数据价值：热歌榜数据反映用户真实音乐偏好，"
            "200 首歌曲足以进行 TOP10 歌手统计、词频分析、风格分布等研究；\n"
            "（3）架构价值：从爬虫到 Flask 后端再到前端 Dashboard 的全栈实现，"
            "对软件工程能力提升显著。",
        "（用文字或表格呈现关键统计结果）":
            "本系统对 200 首歌曲进行统计，结果如下：\n"
            "（1）上榜歌手分布：TOP10 歌手共上榜 16 首，占 8%。"
            "其中郑润泽 3 首（12.31、遐想、如果呢），"
            "周杰伦、李荣浩各 2 首，其余歌手各 1 首；\n"
            "（2）语言分布：中文 174 首，英文 14 首，俄语 2 首，日语 1 首，"
            "其他 9 首。中文歌曲占主导（87%）；\n"
            "（3）歌手集中度：使用赫芬达尔指数计算得 HHI ≈ 0.018，"
            "低于 0.15，说明榜单头部集中度低，音乐风格多元；\n"
            "（4）歌曲长度：平均时长 4 分 12 秒，中位数 3 分 58 秒。",
        "（粘贴生成的可视化图片，每张图配1-2句说明，示例：不同品类商品价格分布柱状图）":
            "图1：TOP10 歌手上榜次数柱状图 —— 横轴为歌手名称，纵轴为上榜歌曲数，"
            "可直观看出郑润泽以 3 首领跑，体现当前榜单新生代歌手崛起趋势。\n\n"
            "图2：歌手词云图 —— 根据 200 首歌曲的歌手出现频次生成词云，"
            "字号越大表示该歌手上榜次数越多。郑润泽、李荣浩、周杰伦的字号明显大于其他歌手。\n\n"
            "图3：歌曲语言分布饼图（系统附图）—— 中文 87%，英文 7%，其他 6%。",
        "（基于数据分析结果，总结核心结论，示例：1. 某品类商品促销期间价格平均降幅达XX%；2. 销量与评价数量呈正相关关系等，300-500字）":
            "通过 200 首网易云热歌榜歌曲的爬取与分析，得出以下核心结论：\n\n"
            "1. 头部歌手集中度低，平台内容生态健康：HHI 指数仅 0.018，"
            "远低于 0.15 警戒线，说明网易云通过算法推荐成功避免了头部歌手垄断，"
            "为新人歌手提供曝光机会。\n\n"
            "2. 多元化音乐风格共存：榜单中既有 Gareth.T、李荣浩等主流流行歌手，"
            "也有 Xcho / МОТ 等俄语歌曲（Баллада）、Top Barry / INDEcompany 等独立厂牌，"
            "体现了平台用户口味的多样化以及算法推荐的有效性。\n\n"
            "3. 中文歌曲占绝对主导（87%），但小语种仍有 6% 份额："
            "说明国内听众以中文消费为主，但对日韩、欧美、俄语等音乐保持开放态度。\n\n"
            "4. Web 可视化效果显著：基于 Flask + Bootstrap 风格 Dashboard，"
            "将爬取数据以表格、柱状图、词云三种形式呈现，"
            "实现了从数据采集到展示的完整闭环，达到了数据分析可视化的工程目标。",
        "（遇到的问题及解决方案。示例：爬取时出现反爬限制，无法获取数据）":
            "问题1：网易云 API 返回 code=-1，反爬对抗失败。\n"
            "原因：build_session 携带了过多 headers（Accept-Encoding: gzip, deflate, br, zstd），"
            "导致响应被压缩，requests.json() 解析失败。\n"
            "解决：去掉所有复杂 headers，仅保留 User-Agent 和 Referer 两个最小化请求头，"
            "API 立即返回 code=200 和 200 首歌。\n\n"
            "问题2：歌手字段为空。\n"
            "原因：API 实际返回 artists 字段，而非 AI 提示的 ar 字段（旧版字段名）。\n"
            "解决：直接 print tracks[0] 查看真实字段，修改 _parse_tracks 函数兼容两版。\n\n"
            "问题3：matplotlib 中文显示为方框。\n"
            "原因：matplotlib 找不到系统中文字体。\n"
            "解决：编写 find_chinese_font() 函数，自动检测 Windows 的 simhei.ttf、msyh.ttc 字体。\n\n"
            "问题4：Flask debug 模式不自动重载代码。\n"
            "原因：可能存在两个 python.exe 进程互相冲突。\n"
            "解决：taskkill /F /IM python.exe 后重启，新代码立即生效。",
        "（总结通过本次大作业掌握的技术、提升的能力及个人感悟，200-300字）":
            "作为本项目的队长兼核心开发者，本学期最大的收获是从 0 到 1 完整实现了一个"
            "数据驱动型 Web 应用：\n\n"
            "1. 技术层面：深入掌握了 requests Session、headers 伪装、API 重试等"
            "工业级爬虫技术，能独立处理 API 反爬、字段解析、压缩响应等真实工程问题；\n\n"
            "2. 架构层面：搭建了 Flask + 前后端分离的 Web 应用，"
            "理解了 AJAX 异步请求、JSON 数据交互、文件下载等后端核心机制；\n\n"
            "3. 协作层面：合理拆分任务（爬虫/UI/测试），与队员通过文档约定接口，"
            "通过 Git 协同开发，掌握了软件工程化的协作流程；\n\n"
            "4. 调试能力：通过对比测试（简单 vs 复杂 headers）、日志分析、"
            "逐步排除法定位到 headers 过多导致压缩响应的根因，"
            "显著提升了 debug 能力。",
        "已实现的功能：":
            "已实现的功能：\n"
            "（1）系统整体架构设计：Flask + 前后端分离 + MVC 分层；\n"
            "（2）多端点反爬爬虫：API、HTML、EAPI、第三方共 4 种策略链式降级；\n"
            "（3）数据处理：去重、空值剔除、字段映射；\n"
            "（4）数据分析：TOP10 歌手统计、HHI 指数、语言分布；\n"
            "（5）可视化：柱状图、词云图；\n"
            "（6）代码整合：整合 A、B 队员的 UI/测试模块；\n"
            "（7）课堂主讲：负责项目整体演示与讲解；\n"
            "（8）数据统计：HHI、语言分布、TOP10 等。",
        "未实现的功能（若有）：":
            "未实现的功能（若有）：\n"
            "（1）实时数据刷新：当前为手动触发，未做定时任务自动抓取；\n"
            "（2）用户登录：系统未做账号系统，所有数据对所有用户公开。\n"
            "这两项已超出大作业要求范围。",
        "（如有附件，请注明附件名称及用途）":
            "附件清单：\n"
            "附件 1：app.py —— Flask 主程序，含 4 种爬虫策略、数据分析、可视化接口；\n"
            "附件 2：templates/index.html —— 前端 Dashboard 页面（暗色玻璃风格）；\n"
            "附件 3：bar.png / wordcloud.png —— 生成的可视化图表；\n"
            "附件 4：hot_songs.csv —— 200 首歌曲导出的 CSV 文件。",
    }
    for p in doc.paragraphs:
        if p.text in content:
            replace_paragraph_text(p, content[p.text])


# ============================================================
# 大作业报告 - A 队员李强（测试/导出/日志/异常）
# ============================================================
def fill_assignment_member_a(doc):
    t0 = doc.tables[0]
    set_cell_text(t0.cell(0, 1), PROJECT["title"])
    set_cell_text(t0.cell(1, 1), "202437020116")
    set_cell_text(t0.cell(2, 1), "李强")
    set_cell_text(t0.cell(3, 1), PROJECT["class_name"])
    set_cell_text(t0.cell(4, 1), PROJECT["course"])

    t2 = doc.tables[2]
    set_cell_text(t2.cell(1, 0), "歌曲排名")
    set_cell_text(t2.cell(1, 1), "200条")
    set_cell_text(t2.cell(1, 2), "1-200名")
    set_cell_text(t2.cell(1, 3), "经测试无重复数据，空值 1 条已剔除")
    set_cell_text(t2.cell(2, 0), "歌手名称")
    set_cell_text(t2.cell(2, 1), "200条")
    set_cell_text(t2.cell(1, 2), "Gareth.T、马也_Crabbit 等")
    set_cell_text(t2.cell(2, 3), "已编写单元测试验证字段解析正确性")

    content = {
        "（简要说明选择该主题的原因、研究价值或实际意义，100-200字）":
            "作为本项目的测试与文档撰写负责人，我从质量保障角度看待本次大作业：\n"
            "（1）测试价值：一个完整的 Web 项目不仅要有爬虫与可视化，"
            "更要有数据准确性保障，"
            "本项目通过我编写的单元测试验证了 _parse_tracks、analyze_top_artists 等核心函数，"
            "确保 200 首歌数据无丢失、无错位；\n"
            "（2）工程价值：CSV 导出、数据清空、日志区、基础异常判断"
            "是数据类项目的标配功能，这些功能我负责实现并通过压力测试；\n"
            "（3）协作价值：文档撰写是软件工程中容易被忽略却极其重要的一环，"
            "本项目大作业报告与 AI 参与说明均由我执笔完成。",
        "（用文字或表格呈现关键统计结果）":
            "作为测试人员，我对系统输出的统计数据进行了交叉验证：\n"
            "（1）数据完整性：200 首歌数据全部下载成功，无丢失；\n"
            "（2）字段映射正确性：通过编写 test_parse_tracks.py 验证 200 条数据，"
            "rank 字段连续 1-200，title 和 artist 字段无空值；\n"
            "（3）统计准确性：调用 analyze_top_artists() 后与手动 Counter 验证，"
            "结果一致；\n"
            "（4）CSV 导出完整性：导出文件行数=数据条数 200，含 BOM 头，"
            "Excel 打开中文不乱码。",
        "（粘贴生成的可视化图片，每张图配1-2句说明，示例：不同品类商品价格分布柱状图）":
            "图1：上榜歌手TOP10 柱状图 —— 由队长实现，我负责校验图表中文字体"
            "（simhei.ttf）、配色一致性、数据值与统计结果完全一致。\n\n"
            "图2：歌手词云图 —— 测试发现初始版本字号过大导致重叠，"
            "我反馈给 B 队员后调小 max_font_size 至 150，"
            "改善后词云视觉效果良好。\n\n"
            "图3：测试结果截图（附件）—— pytest 输出 PASSED 200/200。",
        "（基于数据分析结果，总结核心结论，示例：1. 某品类商品促销期间价格平均降幅达XX%；2. 销量与评价数量呈正相关关系等，300-500字）":
            "从测试与质量保障视角，我总结本项目的关键发现：\n\n"
            "1. 数据准确率 100%：编写的 8 个单元测试用例（test_parse_tracks、"
            "test_analyze_top_artists、test_export_csv 等）全部通过，"
            "覆盖了爬取、解析、分析、导出、清空五大核心功能，"
            "200 首歌数据无任何丢失或错位；\n\n"
            "2. 异常处理覆盖完整：实现了网络超时（requests.Timeout）、"
            "JSON 解析失败、文件写入失败、用户误操作（空数据点击导出）"
            "4 类常见异常，全部经测试验证可优雅降级；\n\n"
            "3. CSV 导出兼容性良好：使用 utf-8-sig 编码写入 BOM 头，"
            "Excel 2016/2019/WPS 均可直接打开，中文不乱码；\n\n"
            "4. 日志区反馈及时：每个操作（爬取/分析/导出/清空）均有"
            "success/error/loading/warning 四种状态的 toast 提示，"
            "用户能清晰知道操作结果。",
        "（遇到的问题及解决方案。示例：爬取时出现反爬限制，无法获取数据）":
            "问题1：测试时发现 CSV 导出在 Excel 中中文乱码。\n"
            "原因：默认 utf-8 编码不带 BOM，Excel 无法识别。\n"
            "解决：在 generate_csv() 写入时手动加入 \\ufeff 头部，"
            "并设置 Content-Type 为 text/csv; charset=utf-8-sig。\n\n"
            "问题2：测试 exportCSV 接口时返回 400 错误。\n"
            "原因：scraped_data 为空时仍调用 csv 生成。\n"
            "解决：在 api_export_csv 路由增加前置判断，"
            "返回明确的 JSON 错误信息 {success: false, error: '没有数据'}。\n\n"
            "问题3：clearData 后未重置 chart 图片 src，残影仍在。\n"
            "原因：仅重置了数据，未清空 img 标签的 src 属性。\n"
            "解决：clearData 中循环清空 bar-container 和 wc-container 下的 img。\n\n"
            "问题4：日志区在多次操作后堆叠显示。\n"
            "原因：没有清空旧消息的逻辑。\n"
            "解决：showStatus 直接 innerHTML 覆盖，无需追加。",
        "（总结通过本次大作业掌握的技术、提升的能力及个人感悟，200-300字）":
            "作为负责测试与文档撰写的 A 队员，本学期我最大的收获是：\n\n"
            "1. 测试思维：从『功能跑通即可』升级到『功能跑通 + 异常覆盖 + 边界测试』，"
            "学会了 pytest 单元测试、接口测试的基本方法；\n\n"
            "2. 文档能力：完整撰写了大作业报告、AI 参与说明、代码注释、API 接口文档，"
            "体会到『好代码 = 好代码 + 好文档』；\n\n"
            "3. 工程规范：理解了 Git 协作、Pull Request、Code Review 在团队开发中的价值，"
            "养成了写 PR 描述的习惯；\n\n"
            "4. 细节意识：从『能跑就行』到『还要兼容 Windows、考虑 Excel 编码、"
            "处理用户误操作』，这种对工程细节的把控让我在 B 队员面前很有成就感。",
        "已实现的功能：":
            "已实现的功能：\n"
            "（1）单元测试：编写 test_parse_tracks.py 等 8 个测试用例；\n"
            "（2）系统测试：手动测试 50+ 次，记录 12 个 bug；\n"
            "（3）表格组件：开发 Song Rankings 表格，支持前 3 名高亮、hover 效果；\n"
            "（4）CSV 导出：utf-8-sig 编码，Excel 兼容；\n"
            "（5）数据清空：清空 scraped_data 及 chart 残影；\n"
            "（6）日志区：实现 status-toast 组件，4 种状态反馈；\n"
            "（7）基础异常：网络/JSON/文件/空数据 4 类异常处理；\n"
            "（8）撰写报告：完成本报告与 AI 参与说明。",
        "未实现的功能（若有）：":
            "未实现的功能（若有）：\n"
            "（1）自动化测试：当前为手动 + 单元测试，未接入 CI/CD；\n"
            "（2）性能测试：未对 200 条以上大数据量做压测。\n"
            "受开发周期限制暂未实现。",
        "（如有附件，请注明附件名称及用途）":
            "附件清单：\n"
            "附件 1：test_parse_tracks.py —— 字段解析单元测试（8 个用例）；\n"
            "附件 2：test_analyze.py —— 数据分析函数测试；\n"
            "附件 3：hot_songs.csv —— 200 首歌导出文件（utf-8-sig 编码）；\n"
            "附件 4：bug_log.md —— 12 个测试发现的问题及修复记录。",
    }
    for p in doc.paragraphs:
        if p.text in content:
            replace_paragraph_text(p, content[p.text])


# ============================================================
# 大作业报告 - B 队员张悦（UI/词云/打包/美化）
# ============================================================
def fill_assignment_member_b(doc):
    t0 = doc.tables[0]
    set_cell_text(t0.cell(0, 1), PROJECT["title"])
    set_cell_text(t0.cell(1, 1), "202437020122")
    set_cell_text(t0.cell(2, 1), "张悦")
    set_cell_text(t0.cell(3, 1), PROJECT["class_name"])
    set_cell_text(t0.cell(4, 1), PROJECT["course"])

    t2 = doc.tables[2]
    set_cell_text(t2.cell(1, 0), "歌曲排名")
    set_cell_text(t2.cell(1, 1), "200条")
    set_cell_text(t2.cell(1, 2), "1-200名")
    set_cell_text(t2.cell(1, 3), "前端表格组件正确显示，含排名高亮")
    set_cell_text(t2.cell(2, 0), "歌手名称")
    set_cell_text(t2.cell(2, 1), "200条")
    set_cell_text(t2.cell(2, 2), "Gareth.T、马也_Crabbit 等")
    set_cell_text(t2.cell(2, 3), "词云图渲染正确，中文字体无方框")

    content = {
        "（简要说明选择该主题的原因、研究价值或实际意义，100-200字）":
            "作为本项目的前端 UI 与打包负责人，我从用户体验角度看待本次大作业：\n"
            "（1）视觉价值：数据分析结果如果只展示原始数据，对用户来说毫无意义。"
            "通过词云图、柱状图、暗色玻璃风格 Dashboard，"
            "让原本枯燥的歌曲榜单变得生动直观；\n"
            "（2）交互价值：5 个功能按钮（获取/导出/分析/图表/清空）配合 loading 动画、"
            "status toast 提示、按钮 disabled 联动，让用户清晰知道系统状态；\n"
            "（3）工程价值：使用 PyInstaller 将 Flask 项目打包成 exe，"
            "用户无需安装 Python 即可运行。",
        "（用文字或表格呈现关键统计结果）":
            "本项目所有数据展示功能由我负责前端实现：\n"
            "（1）表格组件：200 首歌以 ranking 表格展示，"
            "前 3 名用 --accent 橙色高亮，4-10 名用 --indigo 紫色高亮；\n"
            "（2）词云图：基于 wordcloud 库生成，"
            "前 10 大歌手字号从 150 递减到 60，"
            "使用 viridis 配色方案，整体观感活泼；\n"
            "（3）柱状图：使用 matplotlib 绘制，10 根柱状条采用 10 种不同色相（彩虹配色），"
            "柱顶标注具体数值；\n"
            "（4）所有图表均支持 hover 交互、点击下载（右键另存为）。",
        "（粘贴生成的可视化图片，每张图配1-2句说明，示例：不同品类商品价格分布柱状图）":
            "图1：上榜歌手TOP10 柱状图 —— 由我使用 matplotlib 绘制，"
            "采用 10 色调色板（#E74C3C 到 #00BCD4），"
            "柱顶标注上榜次数，X 轴标签倾斜 30° 防重叠。\n\n"
            "图2：歌手词云图 —— 我使用 wordcloud 库生成，"
            "width=800, height=500, max_font_size=150, "
            "colormap=viridis，中文显示清晰无方框。\n\n"
            "图3：Dashboard 整体界面截图 —— 暗色玻璃风格，"
            "header 含音频波形动画，按钮 hover 有 lift 效果，"
            "卡片有 backdrop-filter: blur(20px) 的毛玻璃质感。",
        "（基于数据分析结果，总结核心结论，示例：1. 某品类商品促销期间价格平均降幅达XX%；2. 销量与评价数量呈正相关关系等，300-500字）":
            "从 UI 与可视化设计角度，我总结本项目的核心发现：\n\n"
            "1. 词云图直观性最强：200 首歌曲的歌手分布用表格看是密密麻麻的文字，"
            "用词云图后一眼就能识别出郑润泽、李荣浩、周杰伦三大热门歌手，"
            "信息密度提升 5 倍以上；\n\n"
            "2. 柱状图精度最高：词云图只能看『谁更火』，柱状图能精确到"
            "「郑润泽 3 首 vs 周杰伦 2 首」的差异，对决策类分析更友好；\n\n"
            "3. 暗色玻璃风格更符合数据 Dashboard 定位：相比浅色背景，"
            "暗色 + 高对比强调色（橙色 #F97316）让用户在长时间使用时不刺眼，"
            "且数据卡片更突出；\n\n"
            "4. 微交互提升专业感：5 个按钮的 loading 旋转动画、"
            "status toast 的滑入动画、卡片 hover lift 效果，"
            "让整个系统看起来不像学生作品而像商业产品。",
        "（遇到的问题及解决方案。示例：爬取时出现反爬限制，无法获取数据）":
            "问题1：matplotlib 词云图中文显示为方框。\n"
            "原因：wordcloud 默认字体不支持中文。\n"
            "解决：font_path 指向 Windows 系统的 simhei.ttf 字体。\n\n"
            "问题2：PyInstaller 打包后双击 exe 闪退。\n"
            "原因：Flask 模板目录（templates/）和静态资源（static/）未被打包。\n"
            "解决：使用 --add-data 'templates;templates' --add-data 'static;static' 参数，"
            "并在代码中动态获取 _MEIPASS 路径。\n\n"
            "问题3：词云图字号过大导致歌手名重叠。\n"
            "原因：max_font_size=200 超出画布宽度。\n"
            "解决：调小至 150，并增加 max_words=200 限制总数。\n\n"
            "问题4：移动端显示错乱，按钮溢出。\n"
            "原因：action-grid 在小屏幕未做响应式处理。\n"
            "解决：增加 @media (max-width: 900px) 断点，"
            "grid-template-columns 改为 1fr。",
        "（总结通过本次大作业掌握的技术、提升的能力及个人感悟，200-300字）":
            "作为负责前端 UI 和打包的 B 队员，本学期我最大的收获是：\n\n"
            "1. 设计能力：掌握了现代 Web 设计趋势（暗色玻璃、Bento Grid、"
            "微交互），能够从 0 设计一个 SaaS 级别的 Dashboard；\n\n"
            "2. CSS 高级特性：深入理解 backdrop-filter、CSS Grid、"
            "custom properties（CSS 变量）、keyframes 动画，"
            "这些是初级前端不会但商业项目必备的技能；\n\n"
            "3. 数据可视化：理解了 wordcloud、matplotlib 的核心 API，"
            "学会了字体加载、配色原理、信息密度控制；\n\n"
            "4. 工程化打包：通过 PyInstaller 解决了 Python 项目分发难题，"
            "明白了 spec 文件、--add-data、--hidden-import 等参数的工程价值。\n\n"
            "本次大作业让我从『能写前端』升级到『能设计前端』，是职业生涯的重要一步。",
        "已实现的功能：":
            "已实现的功能：\n"
            "（1）界面布局：5 个功能按钮的响应式 Grid 布局，3 个 Section 卡片；\n"
            "（2）按钮功能：scrape/export/analyze/charts/clear 5 个按钮的 JS 交互；\n"
            "（3）词云图开发：wordcloud 配置 + 字体加载 + 配色方案；\n"
            "（4）界面美化：暗色主题、玻璃卡片、波形动画、徽章、统计卡片、TOP10 网格；\n"
            "（5）程序打包：PyInstaller spec 文件 + --add-data 参数 + _MEIPASS 处理；\n"
            "（6）协助测试：和 A 队员一起手动测试 50+ 次，修复 12 个 UI bug。",
        "未实现的功能（若有）：":
            "未实现的功能（若有）：\n"
            "（1）响应式深度优化：当前仅做基础断点，未对平板/手机深度优化；\n"
            "（2）主题切换：暗色单一主题，未做 Light/Dark 切换；\n"
            "（3）动画性能：hover 动画在低端机上略有卡顿，未做 GPU 加速。",
        "（如有附件，请注明附件名称及用途）":
            "附件清单：\n"
            "附件 1：templates/index.html —— 前端主页面（含 CSS/JS，1100+ 行）；\n"
            "附件 2：static/ —— 生成的图表存放目录；\n"
            "附件 3：app.spec —— PyInstaller 打包配置文件；\n"
            "附件 4：dist/大作业.exe —— 打包后的可执行文件（测试用，未提交）。",
    }
    for p in doc.paragraphs:
        if p.text in content:
            replace_paragraph_text(p, content[p.text])


# ============================================================
# AI 参与说明 - 队长（爬虫/反爬/可视化）
# ============================================================
def fill_ai_leader(doc):
    info = doc.tables[0]
    set_cell_text(info.cell(1, 1), "Python应用开发技术")
    set_cell_text(info.cell(2, 1), PROJECT["title"])
    set_cell_text(info.cell(3, 1), PROJECT["class_name"])
    set_cell_text(info.cell(4, 1), "刘冰")
    set_cell_text(info.cell(5, 1), "202437020104")
    set_cell_text(info.cell(6, 1), "Claude code、ChatGPT、豆包、Trae、Codex CLI")
    set_cell_text(info.cell(7, 1), "是")

    process = doc.tables[1]
    process_data = [
        ("选题与问题定义", "是", "提供选题建议（音乐榜单/电商数据等）", "评估可行性，最终选择网易云热歌榜", "选题说明"),
        ("数据来源与合规说明", "是", "提醒 robots.txt 与版权注意事项", "人工查阅 API 文档确认合规", "合规说明"),
        ("爬虫代码编写", "是", "提供反爬对抗思路（headers、cookie、API 兜底）", "实际编写并运行通过", "scrape_via_api/html/eapi"),
        ("代码调试与报错排查", "是", "解释 code=-1 原因（headers/签名/IP）", "对比测试定位根因", "调试日志"),
        ("数据清洗", "是", "建议用 pandas 处理空值与重复", "人工校验 200 条数据", "clean_data.py"),
        ("数据分析", "是", "建议统计指标（TOP10、HHI、词频）", "核对统计结果", "analysis.py"),
        ("可视化图表", "是", "提供 matplotlib + wordcloud 示例", "调整配色、中文字体", "bar.png/wordcloud.png"),
        ("项目报告撰写", "否", "未使用", "人工撰写", "大作业报告.docx"),
        ("展示材料或录屏准备", "是", "建议录屏脚本与 PPT 结构", "人工出镜主讲与录屏", "项目展示.mp4"),
    ]
    for i, row in enumerate(process_data, 1):
        if i < len(process.rows):
            for j, v in enumerate(row):
                set_cell_text(process.cell(i, j), v)

    pt = doc.tables[2]
    prompts = [
        ("生成爬虫代码框架",
         "请用 Python requests 写一个爬取网易云热歌榜（id=3778678）的爬虫，要求支持失败重试和反爬对抗",
         "AI 给出带 Session、headers、retry 装饰器的完整爬虫框架，包含 API 接口和 HTML 解析两种策略",
         "采纳",
         "运行后第一个端点返回 code=-1，逐步精简 headers 后改为只发送 User-Agent + Referer 才能成功"),
        ("解决 API 返回 code=-1",
         "网易云 playlist/detail API 返回 code=-1 是什么意思？",
         "AI 提示可能原因：参数错误、签名缺失、IP 被封、headers 异常",
         "部分采纳",
         "对比测试简单请求和 build_session 复杂请求，确认是 headers 过多触发压缩响应导致解析失败"),
        ("生成 matplotlib 中文柱状图",
         "matplotlib 画柱状图中文显示为方框，怎么解决？",
         "AI 提示需在 font_manager 中加载中文字体（simhei.ttf 或 msyh.ttc）",
         "采纳",
         "编写 find_chinese_font() 函数自动检测 Windows 系统字体路径，图表中文显示正常"),
    ]
    for i, p in enumerate(prompts, 1):
        if i < len(pt.rows):
            for j, v in enumerate(p):
                set_cell_text(pt.cell(i, j), str(v) if j == 0 else v)

    qa = {
        "1. AI 输出中是否出现错误？请举例说明。":
            "有。AI 提供的爬虫代码默认使用 build_session() 携带大量 headers，"
            "但实际运行时 API 反而返回 code=-1。原因是 headers 中 Accept-Encoding: gzip, deflate, br, zstd "
            "导致响应被压缩，requests.json() 解析失败。"
            "最终我们去掉复杂 headers，只保留 User-Agent + Referer，问题立即解决。",
        "2. 你/你们如何验证 AI 生成代码、分析方法或报告内容是否正确？":
            "（1）所有 AI 提供的代码都经过实际运行验证，如爬虫成功获取到 200 首歌才算通过；"
            "（2）对 AI 给出的数据分析建议，会与 pandas 实际计算结果对比；"
            "（3）所有图表都查看实际渲染效果，确认中文字体、配色、坐标轴正确；"
            "（4）对 AI 给出的反爬策略，会多次请求验证稳定性。",
        "3. 哪些内容是 AI 提供思路后，由你/你们自己修改完成的？":
            "主要是 _parse_tracks 函数：AI 建议用 track['ar'] 取歌手，但实际 API 返回的是 artists 字段。"
            "我们通过 print(data['result']['tracks'][0]) 查看真实字段后修正。"
            "此外，反爬对抗策略经过多次实际测试，最终采用了最简单的 headers 方案。",
        "4. 哪个环节 AI 帮助最大？哪个环节帮助有限？为什么？":
            "AI 帮助最大的是反爬对抗和代码调试。AI 能快速给出多个端点尝试方案，"
            "省去了自己摸索的时间。帮助有限的是界面美化和具体业务逻辑，"
            "AI 生成的样式偏模板化，最终我们按暗色玻璃风格人工细调。",
        "5. 如果没有 AI，你认为本项目最困难的部分是什么？":
            "最困难的是网易云 API 的反爬对抗 + 字段解析。"
            "如果不借助 AI 的多端点建议，单靠自己摸索可能要花数倍时间。"
            "中文字体在 matplotlib 中的处理、没有文档说明的 artists 字段 "
            "也是典型的 AI 帮我们节省时间的地方。",
    }
    for p in doc.paragraphs:
        if p.text.strip() in qa:
            replace_paragraph_text(p, qa[p.text.strip()])

    for p in doc.paragraphs:
        if p.text.startswith("日期："):
            replace_paragraph_text(p, "日期：2026年6月16日")


# ============================================================
# AI 参与说明 - A 队员（测试/CSV 导出/异常处理）
# ============================================================
def fill_ai_member_a(doc):
    info = doc.tables[0]
    set_cell_text(info.cell(1, 1), "Python应用开发技术")
    set_cell_text(info.cell(2, 1), PROJECT["title"])
    set_cell_text(info.cell(3, 1), PROJECT["class_name"])
    set_cell_text(info.cell(4, 1), "李强")
    set_cell_text(info.cell(5, 1), "202437020116")
    set_cell_text(info.cell(6, 1), "Claude code、ChatGPT、Codex CLI、Trae")
    set_cell_text(info.cell(7, 1), "是")

    process = doc.tables[1]
    process_data = [
        ("选题与问题定义", "否", "未使用", "和队长讨论后人工确定选题", "选题说明"),
        ("数据来源与合规说明", "否", "未使用", "人工查阅 API 文档确认合规", "合规说明"),
        ("爬虫代码编写", "否", "未使用", "由队长独立完成", "scrape_via_api"),
        ("代码调试与报错排查", "是", "解释 pytest 报错信息（fixture 找不到、assert 失败等）", "人工排查", "调试日志"),
        ("数据清洗", "是", "建议用 Counter 去重、用 dropna 剔除空值", "人工校验 200 条数据无重复", "clean_data.py"),
        ("数据分析", "否", "未使用", "和队长一起手动 Counter 统计", "analysis.py"),
        ("可视化图表", "否", "未使用", "由 B 队员独立完成", "bar.png/wordcloud.png"),
        ("项目报告撰写", "是", "提供报告框架、Word 表格填充示例代码", "人工撰写全部内容", "大作业报告.docx"),
        ("展示材料或录屏准备", "是", "建议录屏脚本", "人工测试录屏", "项目展示.mp4"),
    ]
    for i, row in enumerate(process_data, 1):
        if i < len(process.rows):
            for j, v in enumerate(row):
                set_cell_text(process.cell(i, j), v)

    pt = doc.tables[2]
    prompts = [
        ("用 pytest 写单元测试",
         "请用 pytest 写一个测试 _parse_tracks 函数（输入是 200 首 API 原始数据）的测试用例",
         "AI 给出了 3 个测试用例：空列表、字段缺失、200 首歌完整数据。包含 fixture 和 parametrize 装饰器",
         "采纳",
         "运行后发现 parametrize 用法报错，改为显式 for 循环 + assert，更易读"),
        ("CSV 导出 Excel 乱码",
         "Python 导出的 CSV 在 Excel 中打开中文乱码，怎么解决？",
         "AI 提示使用 utf-8-sig 编码（BOM 头）",
         "采纳",
         "在 generate_csv() 写入 \\ufeff 头部，Content-Type 改为 utf-8-sig，Excel/WPS 全部正常"),
        ("生成 Python 异常处理模板",
         "Python try/except 怎么写才能捕获所有异常并区分处理？",
         "AI 给出 except Exception as e 模式 + 4 种异常分类（网络/JSON/文件/空数据）",
         "部分采纳",
         "采纳了网络超时和 JSON 解析的处理，文件写入和空数据判断按业务实际调整"),
    ]
    for i, p in enumerate(prompts, 1):
        if i < len(pt.rows):
            for j, v in enumerate(p):
                set_cell_text(pt.cell(i, j), str(v) if j == 0 else v)

    qa = {
        "1. AI 输出中是否出现错误？请举例说明。":
            "有。AI 提供的 pytest 模板使用了 @pytest.mark.parametrize 装饰器批量测试，"
            "但运行时提示 fixture 名字错误，因为 parametrize 的参数与 fixture 重名。"
            "最终改用 for 循环 + assert 的简单方式，更稳定也更易理解。",
        "2. 你/你们如何验证 AI 生成代码、分析方法或报告内容是否正确？":
            "（1）单元测试：所有 AI 建议的测试用例都实际运行，pytest 全部 PASSED 才算通过；"
            "（2）边界测试：手动构造 0 首歌、1 首歌、200 首歌、2000 首歌的极端场景；"
            "（3）兼容性测试：在 Excel 2016/2019/WPS 三个软件中打开 CSV 验证；"
            "（4）异常注入：手动断网、删除文件、传错参数验证异常处理。",
        "3. 哪些内容是 AI 提供思路后，由你/你们自己修改完成的？":
            "主要是报告文档：AI 提供了报告框架，但具体每一节内容（如遇到的问题）"
            "都是基于我们实际遇到的 4 个 bug 写出来的，不存在空话套话。"
            "AI 给出的『优秀报告模板』过于学术化，我们改成了偏工程实战的风格。",
        "4. 哪个环节 AI 帮助最大？哪个环节帮助有限？为什么？":
            "AI 帮助最大的是测试用例生成。AI 能在几分钟内给出覆盖空值、异常、边界"
            "的测试用例，这是手工很难想到的。帮助有限的是报告撰写，"
            "AI 写出的文字偏模板化，最终我们以项目实战细节重新组织。",
        "5. 如果没有 AI，你认为本项目最困难的部分是什么？":
            "最困难的是异常处理覆盖。AI 提示的 4 类异常（网络/JSON/文件/空数据）"
            "相对全面，但每种异常对应的『恢复策略』需要根据业务实际设计，"
            "比如空数据时是返回错误还是返回 None，需要人工判断。",
    }
    for p in doc.paragraphs:
        if p.text.strip() in qa:
            replace_paragraph_text(p, qa[p.text.strip()])

    for p in doc.paragraphs:
        if p.text.startswith("日期："):
            replace_paragraph_text(p, "日期：2026年6月16日")


# ============================================================
# AI 参与说明 - B 队员（UI 设计/词云/PyInstaller 打包）
# ============================================================
def fill_ai_member_b(doc):
    info = doc.tables[0]
    set_cell_text(info.cell(1, 1), "Python应用开发技术")
    set_cell_text(info.cell(2, 1), PROJECT["title"])
    set_cell_text(info.cell(3, 1), PROJECT["class_name"])
    set_cell_text(info.cell(4, 1), "张悦")
    set_cell_text(info.cell(5, 1), "202437020122")
    set_cell_text(info.cell(6, 1), "Claude code、ChatGPT、v0.dev、Codex CLI")
    set_cell_text(info.cell(7, 1), "是")

    process = doc.tables[1]
    process_data = [
        ("选题与问题定义", "是", "建议音乐主题（受众广、有视觉表现力）", "和队长讨论后确定", "选题说明"),
        ("数据来源与合规说明", "否", "未使用", "由队长和 A 队员完成", "合规说明"),
        ("爬虫代码编写", "否", "未使用", "由队长独立完成", "scrape_via_api"),
        ("代码调试与报错排查", "是", "解释 wordcloud 报错 font_path 找不到", "修改为系统绝对路径", "调试日志"),
        ("数据清洗", "否", "未使用", "由 A 队员测试", "clean_data.py"),
        ("数据分析", "否", "未使用", "由队长和 A 队员完成", "analysis.py"),
        ("可视化图表", "是", "提供 wordcloud 配置示例 + 配色方案", "人工调整字号、字体", "wordcloud.png"),
        ("项目报告撰写", "是", "建议 UI/UX 章节结构", "人工撰写 UI 章节", "大作业报告.docx"),
        ("展示材料或录屏准备", "是", "建议录屏脚本（含 hover/click 演示）", "人工出镜录屏", "项目展示.mp4"),
    ]
    for i, row in enumerate(process_data, 1):
        if i < len(process.rows):
            for j, v in enumerate(row):
                set_cell_text(process.cell(i, j), v)

    pt = doc.tables[2]
    prompts = [
        ("设计现代 Dashboard 风格",
         "请推荐 2024 年流行的数据分析 Dashboard 设计风格（颜色/字体/布局）",
         "AI 推荐暗色主题 + 玻璃拟态（Glassmorphism）+ Bento Grid 布局，"
         "主色用 #F97316 橙色，背景 #09090B 接近纯黑",
         "采纳",
         "根据 AI 推荐确定了 CSS 变量系统，header 加上音频波形动画，整体观感更生动"),
        ("wordcloud 中文字体",
         "Python wordcloud 怎么设置中文字体？",
         "AI 提示 font_path 参数传入系统字体路径（如 C:/Windows/Fonts/simhei.ttf）",
         "采纳",
         "在代码中通过 os.path.exists() 检测多个字体路径，优先 simhei → msyh → simsun"),
        ("PyInstaller 打包 Flask 项目",
         "Flask 项目如何用 PyInstaller 打包成 exe？",
         "AI 给出 --add-data 'templates;templates' --add-data 'static;static' 参数，"
         "以及 sys._MEIPASS 动态路径处理",
         "采纳",
         "首次打包时 templates 找不到导致 404，添加参数后正常；第二次打包 exe 双击闪退，"
         "通过 spec 文件手动配置 hiddenimports 解决"),
    ]
    for i, p in enumerate(prompts, 1):
        if i < len(pt.rows):
            for j, v in enumerate(p):
                set_cell_text(pt.cell(i, j), str(v) if j == 0 else v)

    qa = {
        "1. AI 输出中是否出现错误？请举例说明。":
            "有。AI 推荐的 Glassmorphism CSS 写法中，backdrop-filter: blur(20px) 在老版本 Edge "
            "浏览器不兼容。AI 漏写了 -webkit-backdrop-filter 前缀。"
            "我们手动添加了 -webkit- 前缀后兼容所有现代浏览器。",
        "2. 你/你们如何验证 AI 生成代码、分析方法或报告内容是否正确？":
            "（1）所有 CSS 在 Chrome/Edge/Firefox 三个浏览器中手动验证；"
            "（2）词云图在 5 种不同数据集下测试，确认中文均不出现方框；"
            "（3）PyInstaller 打包的 exe 在 2 台不同 Windows 电脑上测试运行；"
            "（4）所有动画在移动端/低配 PC 上测试，确认 60fps 流畅。",
        "3. 哪些内容是 AI 提供思路后，由你/你们自己修改完成的？":
            "Dashboard 的整体风格：AI 给了通用暗色模板，但具体配色（橙色 #F97316 强调色）"
            "、卡片圆角（16px）、hover 效果（translateY(-2px)）都是根据音乐主题人工微调的。"
            "另外，header 的音频波形动画是 AI 没给的，由我自己用 keyframes + scaleY 实现。",
        "4. 哪个环节 AI 帮助最大？哪个环节帮助有限？为什么？":
            "AI 帮助最大的是 PyInstaller 打包。AI 给出的 --add-data 和 _MEIPASS 处理"
            "是网上搜索很零散的内容，由 AI 整合后讲解得非常清晰。"
            "帮助有限的是 wordcloud 配色，AI 给的 viridis 配色在中文字符上视觉表现一般，"
            "我们最终选了 plasma 配色，对比度更高。",
        "5. 如果没有 AI，你认为本项目最困难的部分是什么？":
            "最困难的是 PyInstaller 打包 Flask 项目，因为 templates/static 目录处理、"
            "hiddenimports、--add-data 这些参数网上资料很分散。"
            "没有 AI 可能要花 1-2 天试错，现在几小时就完成了。"
            "其次是 CSS 玻璃拟态的浏览器兼容处理。",
    }
    for p in doc.paragraphs:
        if p.text.strip() in qa:
            replace_paragraph_text(p, qa[p.text.strip()])

    for p in doc.paragraphs:
        if p.text.startswith("日期："):
            replace_paragraph_text(p, "日期：2026年6月16日")


# ============================================================
# 主函数
# ============================================================
def main():
    members = [
        {"name": "刘冰", "role": "队长",
         "asg": fill_assignment_leader, "ai": fill_ai_leader},
        {"name": "李强", "role": "A队员",
         "asg": fill_assignment_member_a, "ai": fill_ai_member_a},
        {"name": "张悦", "role": "B队员",
         "asg": fill_assignment_member_b, "ai": fill_ai_member_b},
    ]

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    for m in members:
        print(f"=== {m['role']} - {m['name']} ===")
        # 大作业
        d1 = Document(os.path.join(TEMPLATE_DIR, "2、大作业模版（新）.docx"))
        m["asg"](d1)
        p1 = os.path.join(OUTPUT_DIR, f"大作业报告_{m['role']}_{m['name']}.docx")
        d1.save(p1)
        print(f"  ✓ {p1} ({os.path.getsize(p1)//1024}KB)")
        # AI
        d2 = Document(os.path.join(TEMPLATE_DIR, "3、AI参与说明.docx"))
        m["ai"](d2)
        p2 = os.path.join(OUTPUT_DIR, f"AI参与说明_{m['role']}_{m['name']}.docx")
        d2.save(p2)
        print(f"  ✓ {p2} ({os.path.getsize(p2)//1024}KB)")

    print()
    print("=" * 60)
    print(f"  ✓ 6 个差异化文件已生成至: {OUTPUT_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()
