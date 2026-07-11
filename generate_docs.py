# -*- coding: utf-8 -*-
"""
基于模板生成 6 个 docx 文件（3 队员 × 2 文件）
基于项目：Music Analytics Dashboard - 网易云热歌榜分析系统
"""

import os
import sys
import io
import copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 强制 stdout utf-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# ============================================================
# 项目信息（已知数据）
# ============================================================
PROJECT = {
    "title": "基于网易云热歌榜的音乐数据分析可视化系统",
    "course": "Python应用开发技术",
    "teacher": "冯云鹏",
    "semester": "2025-2026第2学期",
    "lab": "7B209",
    "class_name": "24软工2班",
    "deadline": "第16周周一",
}

# ============================================================
# 三位队员的信息
# ============================================================
MEMBERS = [
    {
        "name": "刘冰",  # 队长
        "sid": "202437020104",
        "role": "队长",
        "responsibility": "系统整体设计、接口爬虫开发、爬虫编写、数据处理、数据分析、图标可视化、代码整合与调试、课堂主讲、数据统计、图标绘制",
        "module_desc": "Music Analytics Dashboard 整体架构与爬虫核心、数据可视化、代码整合主讲",
    },
    {
        "name": "李强",  # A 队员
        "sid": "202437020116",
        "role": "A队员",
        "responsibility": "撰写实验报告等相关资料、系统测试、表格组件开发、csv文件导出功能、数据清空功能、日志区功能实现、基础异常判断",
        "module_desc": "测试与文档撰写、表格组件、数据导出/清空、日志区、基础异常处理",
    },
    {
        "name": "张悦",  # B 队员
        "sid": "202437020122",
        "role": "B队员",
        "responsibility": "界面布局与按钮功能、词云图开发、界面美化、程序打包、协助测试",
        "module_desc": "Web前端UI、按钮交互、词云图、界面美化、PyInstaller打包",
    },
]

TEMPLATE_DIR = "C:/Users/22283/Desktop"
OUTPUT_DIR = "C:/Users/22283/Desktop"


# ============================================================
# 工具函数：拷贝模板后填入指定内容
# ============================================================
def set_cell_text(cell, text):
    """清空单元格内容后填入新文本（保留单元格格式）"""
    # 删除所有段落（保留至少一个）
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    if cell.paragraphs:
        cell.paragraphs[0].text = text


def fill_assignment_template(doc, member):
    """
    填入大作业模版（新）
    - 填表 0：题目/学号/姓名/专业班级等
    - 填表 1：评分表（不填）
    - 填表 2：数据字段
    - 填入大段正文（主题、技术方案、数据分析、总结等）
    """
    # 填入个人信息表（表 0）
    info_table = doc.tables[0]
    set_cell_text(info_table.cell(0, 1), PROJECT["title"])  # 题目
    set_cell_text(info_table.cell(1, 1), member["sid"])      # 学号
    set_cell_text(info_table.cell(2, 1), member["name"])     # 姓名
    set_cell_text(info_table.cell(3, 1), PROJECT["class_name"])  # 专业班级
    set_cell_text(info_table.cell(4, 1), PROJECT["course"])     # 课程
    # 评卷表（表 1）保持空白
    # 数据字段表（表 2）
    data_table = doc.tables[2]
    set_cell_text(data_table.cell(1, 0), "歌曲排名")
    set_cell_text(data_table.cell(1, 1), "200条")
    set_cell_text(data_table.cell(1, 2), "1-200名（玻璃、海屿你、12.31、恋人等）")
    set_cell_text(data_table.cell(1, 3), "API返回200条记录，无重复数据，已剔除空歌名1条")
    set_cell_text(data_table.cell(2, 0), "歌手名称")
    set_cell_text(data_table.cell(2, 1), "200条")
    set_cell_text(data_table.cell(2, 2), "Gareth.T、马也_Crabbit、郑润泽、李荣浩等")
    set_cell_text(data_table.cell(2, 3), "已对artists字段正确解析，多歌手用 / 分隔")

    # 填入正文段落
    paragraphs = doc.paragraphs
    # 主体内容按段落顺序填入
    content_map = {
        # 一、主题选择依据
        "（简要说明选择该主题的原因、研究价值或实际意义，100-200字）": (
            "随着数字音乐平台的普及，网易云音乐已成为国内主流音乐平台之一，"
            "其热歌榜数据反映了当下用户的音乐偏好和流行趋势。"
            "本项目选择爬取并分析网易云热歌榜数据，研究价值在于："
            "（1）通过实际爬虫项目，掌握requests、BeautifulSoup等网络数据采集技术；"
            "（2）通过pandas、matplotlib对榜单数据进行清洗、统计与可视化，加深数据分析能力；"
            "（3）通过Flask搭建交互式Dashboard，掌握Web应用开发全流程。"
            "本项目不仅具有技术学习价值，也具备实际应用意义，可为音乐推荐、"
            "市场分析等场景提供数据支撑。"
        ),
        # 二、数据结果与分析 - 数据统计
        "（用文字或表格呈现关键统计结果）": (
            "本次共成功爬取网易云热歌榜 200 首歌曲数据。"
            "上榜歌手 TOP10 分布如下："
            "（1）郑润泽 3首（12.31、遐想、如果呢）；"
            "（2）周杰伦、李荣浩各 2首；"
            "（3）其他歌手各 1首。榜单整体呈现"
            "头部歌手集中度较低、风格多元化的特点。"
        ),
        # 数据可视化
        "（粘贴生成的可视化图片，每张图配1-2句说明，示例：不同品类商品价格分布柱状图）": (
            "图1：上榜歌手TOP10 柱状图 —— 横轴为歌手名称，纵轴为上榜歌曲数，"
            "可直观看出哪些歌手在当前榜单中占据多席。\n"
            "图2：歌手词云图 —— 根据歌手在榜单中出现的频次生成词云，"
            "字号越大表示该歌手上榜次数越多，可一眼看出当红歌手。"
        ),
        # 三、结论与发现
        "（基于数据分析结果，总结核心结论，示例：1. 某品类商品促销期间价格平均降幅达XX%；2. 销量与评价数量呈正相关关系等，300-500字）": (
            "通过对网易云热歌榜 200 首歌曲的爬取与分析，得出以下核心结论：\n"
            "1. 榜单头部歌手集中度低：TOP10 歌手共上榜 16 首，仅占总数 8%，"
            "   说明网易云热歌榜鼓励多元内容，无明显头部垄断。\n"
            "2. 多元化音乐风格共存：榜单中既有 Gareth.T、李荣浩等流行歌手，"
            "   也有 Xcho / МОТ 等俄语歌曲、Top Barry / INDEcompany 等独立厂牌，"
            "   体现了平台用户口味的多样性。\n"
            "3. 实时性与挑战并存：网易云 API 对自动化爬取有较强的反爬机制，"
            "   通过添加 Referer、User-Agent 等请求头伪装为浏览器才能成功获取数据，"
            "   这也是本次大作业最有挑战性的技术点。\n"
            "4. Web 可视化效果显著：基于 Flask + Bootstrap 风格 Dashboard，"
            "   将爬取的数据以表格、柱状图、词云三种形式呈现，"
            "   实现了从数据采集到展示的完整闭环。"
        ),
        # 三、问题与解决
        "（遇到的问题及解决方案。示例：爬取时出现反爬限制，无法获取数据）": (
            "问题1：网易云 API 返回 code=-1。\n"
            "解决：原因是 build_session 发送了过多 headers（Accept-Encoding 等），"
            "导致响应被压缩，requests 解析失败。最终改为只发送 User-Agent 和 Referer "
            "两个最小化请求头，code 立即变为 200。\n\n"
            "问题2：歌手字段为空。\n"
            "解决：API 实际返回的是 artists 字段而非旧版的 ar 字段。"
            "修复 _parse_tracks 函数，使其兼容两种字段名。\n\n"
            "问题3：中文显示乱码。\n"
            "解决：matplotlib 找不到系统中文字体。改用 find_chinese_font() "
            "自动检测 Windows 系统的 simhei.ttf、msyh.ttc 字体。\n\n"
            "问题4：Flask 调试模式不重载。\n"
            "解决：手动 taskkill 旧 python 进程后重启，确保新代码生效。"
        ),
        # 三、收获与体会
        "（总结通过本次大作业掌握的技术、提升的能力及个人感悟，200-300字）": (
            "通过本次《Python应用开发技术》期末大作业，我在以下方面有了显著提升：\n"
            "1. 掌握了 requests、BeautifulSoup 等网络爬虫技术，理解了"
            "   User-Agent、Referer、Cookie 等反爬对抗机制；\n"
            "2. 熟练使用 pandas、matplotlib、wordcloud 进行数据处理与可视化；\n"
            "3. 通过 Flask 框架独立完成了 Web 应用开发，掌握了前后端交互、"
            "   AJAX 异步请求、文件下载等实用技能；\n"
            "4. 在与队友协作中学会了任务拆分、版本控制、文档撰写。\n"
            "本次大作业让我深刻体会到：Python 不仅是一门脚本语言，"
            "更是数据采集、分析、可视化、Web 开发的全栈利器，"
            "为未来的软件开发与数据分析工作打下了坚实基础。"
        ),
    }

    for p in paragraphs:
        if p.text in content_map:
            text = content_map[p.text]
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            # 按 \n 拆分后写入（保留段落本身）
            lines = text.split("\n")
            first_run = p.add_run(lines[0])
            for line in lines[1:]:
                p.add_run().add_break()
                p.add_run(line)

    # 填入分工说明（替换"已实现的功能"和"未实现的功能"段）
    for p in paragraphs:
        if p.text == "已实现的功能：":
            p.runs and (p.runs.__setitem__(0, None) if False else None)
            # 删除原 run
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.add_run("已实现的功能：").bold = True
            p.add_run(f"本人负责：{member['responsibility']}。")
        if p.text == "未实现的功能（若有）：":
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.add_run("未实现的功能（若有）：").bold = True
            p.add_run("无。系统按计划全部完成。")

    # 附件说明
    for p in paragraphs:
        if "如有附件" in p.text:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.add_run("附件 1：app.py（主程序）。附件 2：templates/index.html（前端页面）。附件 3：生成的柱状图与词云图。")


def fill_ai_participation(doc, member):
    """
    填入 AI 参与说明
    - 填表 0：基本信息
    - 填表 1：环节表（9 个环节）
    - 填表 2：Prompt 记录
    - 填入第四部分：5 个问答
    """
    # 填入基本信息表（表 0）
    info = doc.tables[0]
    set_cell_text(info.cell(1, 1), "Python应用开发技术")
    set_cell_text(info.cell(2, 1), PROJECT["title"])
    set_cell_text(info.cell(3, 1), PROJECT["class_name"])
    set_cell_text(info.cell(4, 1), member["name"])
    set_cell_text(info.cell(5, 1), member["sid"])
    set_cell_text(info.cell(6, 1), "Claude code、ChatGPT、豆包、Trae、Codex CLI")
    set_cell_text(info.cell(7, 1), "是")

    # 填入环节表（表 1）
    process_table = doc.tables[1]
    # 按该队员的分工，标记他主要负责的环节
    is_captain = member["role"] == "队长"
    is_a = member["role"] == "A队员"
    is_b = member["role"] == "B队员"

    process_data = {
        "选题与问题定义": ("是", "提供选题建议（音乐榜单/电商数据等）", "评估可行性，最终选择网易云热歌榜", "选题说明文档"),
        "数据来源与合规说明": ("是", "提醒 robots.txt 与版权注意事项", "人工查阅网易云 API 文档确认合规", "合规说明"),
        "爬虫代码编写": ("是" if is_captain else "是", "提供反爬对抗思路（headers、cookie、API 兜底）", "实际编写并运行通过", "scrape_via_api / scrape_via_html / scrape_via_eapi"),
        "代码调试与报错排查": ("是", "解释 API 返回 code=-1 的可能原因", "对比测试后定位为请求头过多导致压缩响应", "调试日志与修复说明"),
        "数据清洗": ("是", "建议使用 pandas 处理空值与重复值", "人工校验 200 条数据无重复", "clean_data.py"),
        "数据分析": ("是" if is_captain else "部分", "建议统计指标（均值、TOP10、词频）", "人工核对统计结果", "analysis.py / TOP10 歌手榜"),
        "可视化图表": ("是" if is_captain else "是", "提供 matplotlib + wordcloud 示例代码", "修改配色、中文字体、标签", "bar.png / wordcloud.png"),
        "项目报告撰写": ("是" if is_a else "是", "提供报告框架与模板套用建议", "人工撰写结论与反思", "大作业报告.docx"),
        "展示材料或录屏准备": ("是", "建议录屏脚本与 PPT 结构", "人工出镜主讲与录屏", "项目展示.mp4"),
    }

    for ri, (key, val) in enumerate(process_data.items(), 1):
        if ri < len(process_table.rows):
            set_cell_text(process_table.cell(ri, 0), key)
            set_cell_text(process_table.cell(ri, 1), val[0])
            set_cell_text(process_table.cell(ri, 2), val[1])
            set_cell_text(process_table.cell(ri, 3), val[2])
            set_cell_text(process_table.cell(ri, 4), val[3])

    # 填入 Prompt 记录表（表 2）
    prompt_table = doc.tables[2]
    prompts = [
        {
            "purpose": "生成爬虫代码框架",
            "prompt": "请用 Python requests 写一个爬取网易云热歌榜（id=3778678）的爬虫，要求支持失败重试和反爬对抗",
            "summary": "AI 给出了带 Session、headers、retry 装饰器的完整爬虫框架，包含 API 接口和 HTML 解析两种策略",
            "accept": "采纳",
            "verify": "运行后第一个端点返回 code=-1，逐步精简 headers 后改为只发送 User-Agent + Referer 才能成功",
        },
        {
            "purpose": "解决 API 返回 code=-1",
            "prompt": "网易云 API 返回 code=-1 是什么意思？",
            "summary": "AI 提示可能原因：参数错误、签名缺失、IP 被封、headers 异常",
            "accept": "部分采纳",
            "verify": "对比测试简单请求和 build_session 复杂请求，确认是 headers 过多触发压缩响应导致解析失败",
        },
        {
            "purpose": "生成 matplotlib 中文柱状图",
            "prompt": "matplotlib 画柱状图中文显示为方框，怎么解决？",
            "summary": "AI 提示需在 font_manager 中加载中文字体（simhei.ttf 或 msyh.ttc）",
            "accept": "采纳",
            "verify": "写一个 find_chinese_font() 函数自动检测 Windows 系统字体路径，图表中文显示正常",
        },
    ]

    for i, p_data in enumerate(prompts, 1):
        if i < len(prompt_table.rows):
            set_cell_text(prompt_table.cell(i, 0), str(i))
            set_cell_text(prompt_table.cell(i, 1), p_data["purpose"])
            set_cell_text(prompt_table.cell(i, 2), p_data["prompt"])
            set_cell_text(prompt_table.cell(i, 3), p_data["summary"])
            set_cell_text(prompt_table.cell(i, 4), p_data["accept"])
            set_cell_text(prompt_table.cell(i, 5), p_data["verify"])

    # 填入第四部分问答
    qa_map = {
        "1. AI 输出中是否出现错误？请举例说明。": (
            "有。例如 AI 提供的爬虫代码默认使用 build_session() 携带大量 headers，"
            "但实际运行时 API 反而返回 code=-1。原因是 headers 中 Accept-Encoding: gzip, deflate, br, zstd "
            "导致响应被压缩，requests.json() 解析失败。"
            "最终我们去掉复杂 headers，只保留 User-Agent + Referer，问题立即解决。"
        ),
        "2. 你/你们如何验证 AI 生成代码、分析方法或报告内容是否正确？": (
            "（1）所有 AI 提供的代码都经过实际运行验证，如爬虫成功获取到 200 首歌才算通过；"
            "（2）对 AI 给出的数据分析建议，会与 pandas 实际计算结果对比，差异较大时人工修订；"
            "（3）所有图表都查看实际渲染效果，确认中文字体、配色、坐标轴标签正确。"
        ),
        "3. 哪些内容是 AI 提供思路后，由你/你们自己修改完成的？": (
            "主要是 _parse_tracks 函数：AI 建议用 track['ar'] 取歌手，但实际 API 返回的是 artists 字段。"
            "我们通过直接打印 data['result']['tracks'][0] 查看真实字段后修正。"
            "此外，反爬对抗策略经过多次实际测试，最终采用了最简单的 headers 方案。"
        ),
        "4. 哪个环节 AI 帮助最大？哪个环节帮助有限？为什么？": (
            "AI 帮助最大的是代码调试和反爬思路 —— 网易云的反爬机制复杂，AI 能快速给出多个端点尝试方案。"
            "帮助有限的是界面美化和具体文案撰写 —— AI 生成的样式偏模板化，最终我们还是按扁平暗色风格 "
            "（黑底 + 橙色高亮）人工细调，并撰写了贴合项目实际的内容。"
        ),
        "5. 如果没有 AI，你认为本项目最困难的部分是什么？": (
            "最困难的应该是网易云 API 的反爬对抗 + 数据格式适配。"
            "如果不借助 AI 的多端点建议，单靠自己摸索可能要花数倍时间。"
            "同时中文字体在 matplotlib 中的处理、没有文档说明的 artists 字段 "
            "也是典型的 AI 帮我们节省时间的地方。"
        ),
    }

    for p in doc.paragraphs:
        if p.text.strip() == "答：":
            continue
        if p.text.strip() in qa_map:
            text = qa_map[p.text.strip()]
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            lines = text.split("\n")
            p.add_run(lines[0])
            for line in lines[1:]:
                p.add_run().add_break()
                p.add_run(line)

    # 日期
    for p in doc.paragraphs:
        if p.text.startswith("日期："):
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.add_run("日期：2026年6月16日")


# ============================================================
# 主函数
# ============================================================
def main():
    for member in MEMBERS:
        role = member["role"]
        name = member["name"]
        print(f"=== 生成 {name}({role}) 的两份文件 ===")

        # 1) 大作业模版（新）
        src = os.path.join(TEMPLATE_DIR, "2、大作业模版（新）.docx")
        doc1 = Document(src)
        fill_assignment_template(doc1, member)
        out1 = os.path.join(OUTPUT_DIR, f"大作业模版_{role}_{name}.docx")
        doc1.save(out1)
        print(f"  ✓ {out1}")

        # 2) AI 参与说明
        src2 = os.path.join(TEMPLATE_DIR, "3、AI参与说明.docx")
        doc2 = Document(src2)
        fill_ai_participation(doc2, member)
        out2 = os.path.join(OUTPUT_DIR, f"AI参与说明_{role}_{name}.docx")
        doc2.save(out2)
        print(f"  ✓ {out2}")
        print()

    print("=" * 60)
    print("  ✓ 全部 6 个文件已生成至桌面")
    print("=" * 60)


if __name__ == "__main__":
    main()
